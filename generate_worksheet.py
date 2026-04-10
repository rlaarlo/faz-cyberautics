"""
Generate worksheet_public_places.docx
Topic: Public Places – Identifying Explicit & Implicit Information
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_border(cell, **kwargs):
    """Set borders on a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), kwargs.get(edge, "single"))
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")
        tcBorders.append(element)
    tcPr.append(tcBorders)


def set_row_borders(row):
    for cell in row.cells:
        set_cell_border(cell)


def add_heading(doc, text, level=1, font_size=13, bold=True, center=False):
    para = doc.add_paragraph()
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(font_size)
    return para


def add_body(doc, text, bold=False, italic=False):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    return para


def style_header_cell(cell, text):
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "D9E1F2")
    tcPr.append(shd)


def set_cell_text(cell, text, center=False, bold=False, italic=False):
    cell.text = ""
    para = cell.paragraphs[0]
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Calibri"
    run.font.size = Pt(11)


def main():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # ── Title block ──────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Worksheet: Building Knowledge of the Field")
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(16)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Topic: Public Places")
    r2.bold = True
    r2.font.name = "Calibri"
    r2.font.size = Pt(13)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("Goal: Identifying Explicit & Implicit Information")
    r3.font.name = "Calibri"
    r3.font.size = Pt(13)

    doc.add_paragraph()  # spacer

    # ── Activity 1 ──────────────────────────────────────────────────────────
    add_heading(doc, "Activity 1 – Vocabulary Mapping", font_size=13)
    add_body(
        doc,
        "Instruction: Match each public place with its correct description "
        "and one thing you can do there.",
        italic=True,
    )

    headers1 = ["No.", "Public Place", "Description", "What can you do there?"]
    rows1 = [
        ["1", "Hospital", "a. A place where you can borrow and read books", "( ) Buy medicine"],
        ["2", "Library", "b. A place where sick people are treated by doctors", "( ) Watch a movie"],
        ["3", "Post Office", "c. A place where you can buy stamps and send letters", "( ) Read and borrow books"],
        ["4", "Bank", "d. A place where you can watch films on a big screen", "( ) See a doctor"],
        ["5", "Cinema", "e. A place where you can save or withdraw money", "( ) Send a letter"],
        ["6", "Pharmacy", "f. A place where you can buy medicine", "( ) Save money"],
    ]
    col_widths1 = [Cm(1.2), Cm(3.5), Cm(8.0), Cm(5.0)]

    table1 = doc.add_table(rows=1 + len(rows1), cols=4)
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (cell, w) in enumerate(zip(table1.rows[0].cells, col_widths1)):
        cell.width = w
        style_header_cell(cell, headers1[i])
        set_cell_border(cell)
    for row_data in rows1:
        row = table1.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].width = col_widths1[i]
            set_cell_text(row.cells[i], val, center=(i == 0))
            set_cell_border(row.cells[i])

    doc.add_paragraph()

    # ── Activity 2 ──────────────────────────────────────────────────────────
    add_heading(doc, "Activity 2 – Picture Observation", font_size=13)
    add_body(
        doc,
        "Instruction: Look at the signs below. Write where you can find each sign "
        "and what the sign means (implicit message).",
        italic=True,
    )

    headers2 = ["No.", "Sign Text", "Where can you find it?", "What does it really tell you? (Implicit)"]
    rows2 = [
        ["1", '"No Running"', "_________________", "_________________"],
        ["2", '"Silence, Please"', "_________________", "_________________"],
        ["3", '"Wash Your Hands"', "_________________", "_________________"],
        ["4", '"Show Your Ticket"', "_________________", "_________________"],
        ["5", '"Turn Off Your Phone"', "_________________", "_________________"],
    ]
    col_widths2 = [Cm(1.2), Cm(4.0), Cm(5.5), Cm(7.0)]

    table2 = doc.add_table(rows=1 + len(rows2), cols=4)
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (cell, w) in enumerate(zip(table2.rows[0].cells, col_widths2)):
        cell.width = w
        style_header_cell(cell, headers2[i])
        set_cell_border(cell)
    for row_data in rows2:
        row = table2.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].width = col_widths2[i]
            set_cell_text(row.cells[i], val, center=(i == 0))
            set_cell_border(row.cells[i])

    # Example box
    doc.add_paragraph()
    p_ex = doc.add_paragraph()
    r_ex = p_ex.add_run("Example:")
    r_ex.bold = True
    r_ex.font.name = "Calibri"
    r_ex.font.size = Pt(11)

    example_lines = [
        ('Sign: ', '"Wet Floor"'),
        ('Where: ', 'Hospital / Mall'),
        ('Explicit meaning: ', 'The floor is wet.'),
        (
            'Implicit meaning: ',
            'You should be careful and walk slowly so you don\'t slip.',
        ),
    ]
    for label, content in example_lines:
        p = doc.add_paragraph(style="List Bullet")
        r_label = p.add_run(label)
        r_label.bold = True
        r_label.font.name = "Calibri"
        r_label.font.size = Pt(11)
        r_content = p.add_run(content)
        r_content.font.name = "Calibri"
        r_content.font.size = Pt(11)

    doc.add_paragraph()

    # ── Activity 3 ──────────────────────────────────────────────────────────
    add_heading(doc, "Activity 3 – Read & Identify (Explicit vs. Implicit)", font_size=13)
    add_body(
        doc,
        "Instruction: Read the short text below, then answer the questions. "
        "Write (E) for Explicit and (I) for Implicit.",
        italic=True,
    )

    # Text 1
    add_heading(doc, "Text 1", font_size=11, bold=True)
    p_t1 = doc.add_paragraph()
    r_t1 = p_t1.add_run(
        "\u201cThe City Library is open from Monday to Saturday, 08.00 a.m. \u2013 04.00 p.m. "
        "Visitors must bring their library card. Children under 10 must be accompanied by an adult. "
        "Please keep quiet and do not eat or drink inside the library.\u201d"
    )
    r_t1.italic = True
    r_t1.font.name = "Calibri"
    r_t1.font.size = Pt(11)

    headers3 = ["No.", "Question", "Answer", "E / I"]
    rows3a = [
        ["1", "What time does the library open?", "_________________", "____"],
        ["2", "Can you visit the library on Sunday?", "_________________", "____"],
        ["3", "What must visitors bring?", "_________________", "____"],
        ["4", "Why should children under 10 be with an adult?", "_________________", "____"],
        ["5", "Can you eat a sandwich inside the library?", "_________________", "____"],
        ["6", "Is the library a noisy place?", "_________________", "____"],
    ]
    col_widths3 = [Cm(1.2), Cm(7.5), Cm(5.5), Cm(2.0)]

    table3a = doc.add_table(rows=1 + len(rows3a), cols=4)
    table3a.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (cell, w) in enumerate(zip(table3a.rows[0].cells, col_widths3)):
        cell.width = w
        style_header_cell(cell, headers3[i])
        set_cell_border(cell)
    for row_data in rows3a:
        row = table3a.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].width = col_widths3[i]
            set_cell_text(row.cells[i], val, center=(i in (0, 3)))
            set_cell_border(row.cells[i])

    doc.add_paragraph()

    # Text 2
    add_heading(doc, "Text 2", font_size=11, bold=True)
    p_t2 = doc.add_paragraph()
    r_t2 = p_t2.add_run(
        "\u201cAttention, please. The next train to Surabaya will depart from Platform 3 at 10.15 a.m. "
        "Passengers are reminded to keep their belongings safe. "
        "Please stand behind the yellow line while waiting.\u201d"
    )
    r_t2.italic = True
    r_t2.font.name = "Calibri"
    r_t2.font.size = Pt(11)

    rows3b = [
        ["1", "Where will you hear this announcement?", "_________________", "____"],
        ["2", "What time will the train depart?", "_________________", "____"],
        ["3", "Why should passengers stand behind the yellow line?", "_________________", "____"],
        ["4", "Where is the train going?", "_________________", "____"],
        ["5", "Is it possible that someone might steal your bag here?", "_________________", "____"],
    ]

    table3b = doc.add_table(rows=1 + len(rows3b), cols=4)
    table3b.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (cell, w) in enumerate(zip(table3b.rows[0].cells, col_widths3)):
        cell.width = w
        style_header_cell(cell, headers3[i])
        set_cell_border(cell)
    for row_data in rows3b:
        row = table3b.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].width = col_widths3[i]
            set_cell_text(row.cells[i], val, center=(i in (0, 3)))
            set_cell_border(row.cells[i])

    doc.add_paragraph()

    # ── Activity 4 ──────────────────────────────────────────────────────────
    add_heading(doc, "Activity 4 – Class Discussion", font_size=13)
    add_body(
        doc,
        "Instruction: Discuss these questions with your partner.",
        italic=True,
    )

    discussion_questions = [
        "What public places do you visit most often? Why?",
        "What rules do you usually find in those places?",
        "Why do public places have rules? What would happen without them?",
        "Can you think of a sign that says one thing but means something more?",
    ]
    for i, q in enumerate(discussion_questions, 1):
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(q)
        r.font.name = "Calibri"
        r.font.size = Pt(11)

    doc.add_paragraph()

    # ── Activity 5 ──────────────────────────────────────────────────────────
    add_heading(doc, "Activity 5 – Let's Create!", font_size=13)
    add_body(
        doc,
        "Instruction: Choose ONE public place. Create a short notice/announcement "
        "for that place. Then write:",
        italic=True,
    )

    create_items = [
        ("My Public Place: ", "_________________________________"),
        ("My Notice/Announcement:", ""),
        ("", "_________________________________________________________________"),
        ("", "_________________________________________________________________"),
        ("Explicit message (what it directly says): ", "_________________________________"),
        ("Implicit message (what it really means): ", "_________________________________"),
    ]
    for label, content in create_items:
        p = doc.add_paragraph()
        if label:
            r_label = p.add_run(label)
            r_label.bold = True
            r_label.font.name = "Calibri"
            r_label.font.size = Pt(11)
        r_content = p.add_run(content)
        r_content.font.name = "Calibri"
        r_content.font.size = Pt(11)

    doc.add_paragraph()

    # ── Key Concepts ─────────────────────────────────────────────────────────
    add_heading(doc, "Key Concepts to Remember", font_size=13)

    headers_k = ["Term", "Meaning", "Example"]
    rows_k = [
        [
            "Explicit (Tersurat)",
            "Information that is directly stated in the text.",
            "\u201cThe library opens at 08.00 a.m.\u201d",
        ],
        [
            "Implicit (Tersirat)",
            "Information that is not directly stated but can be understood/inferred.",
            "\u201cPlease keep quiet\u201d \u2192 The library is a place that requires silence; "
            "it is not a place for chatting.",
        ],
    ]
    col_widths_k = [Cm(3.5), Cm(7.5), Cm(6.7)]

    table_k = doc.add_table(rows=1 + len(rows_k), cols=3)
    table_k.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (cell, w) in enumerate(zip(table_k.rows[0].cells, col_widths_k)):
        cell.width = w
        style_header_cell(cell, headers_k[i])
        set_cell_border(cell)
    for row_data in rows_k:
        row = table_k.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].width = col_widths_k[i]
            set_cell_text(row.cells[i], val, bold=(i == 0))
            set_cell_border(row.cells[i])

    # Save
    output_path = "worksheet_public_places.docx"
    doc.save(output_path)
    print(f"Saved: {output_path}")


if __name__ == "__main__":
    main()
